"""
Document import via AI:
- PDF: Render pages as images via PyMuPDF, send to Claude multimodal
- DOCX: Extract text + images directly via python-docx, send to Claude
No external tools (LibreOffice, pandoc) required.
"""

import json
import re
import base64
import asyncio
import io

from anthropic import Anthropic
from config import ANTHROPIC_API_KEY, CLAUDE_MODEL

# Rate limiting
_semaphore = asyncio.Semaphore(2)

_client = None


def _get_client():
    global _client
    if _client is None:
        _client = Anthropic(api_key=ANTHROPIC_API_KEY)
    return _client


def _parse_json_response(text: str) -> list:
    """Extract JSON array from Claude's response, with repair for common issues."""
    text = text.strip()
    if text.startswith("```"):
        lines = text.split("\n")
        lines = lines[1:]
        if lines and lines[-1].strip() == "```":
            lines = lines[:-1]
        text = "\n".join(lines)

    # First attempt: parse as-is
    try:
        return json.loads(text)
    except json.JSONDecodeError:
        pass

    # Repair attempt: fix common Claude JSON issues
    repaired = text

    # Fix truncated output: close open brackets/braces
    if repaired.count("[") > repaired.count("]"):
        # Find last complete object (ending with })
        last_brace = repaired.rfind("}")
        if last_brace != -1:
            repaired = repaired[: last_brace + 1]
            # Remove trailing comma if present
            repaired = repaired.rstrip().rstrip(",")
            repaired += "]"

    # Remove trailing commas before ] or }
    repaired = re.sub(r",\s*([}\]])", r"\1", repaired)

    try:
        return json.loads(repaired)
    except json.JSONDecodeError:
        pass

    # Last resort: find the outermost JSON array with regex
    match = re.search(r"\[.*\]", repaired, re.DOTALL)
    if match:
        try:
            return json.loads(match.group())
        except json.JSONDecodeError:
            pass

    raise ValueError(f"Konnte kein gültiges JSON aus der KI-Antwort extrahieren.")


_DOC_TYPE_DESCRIPTIONS = {
    "multichoice": 'multichoice: Multiple Choice (question_data: {"single": true, "shuffle": true, "answers": [{"text": "...", "fraction": 100, "feedback": "..."}]})',
    "truefalse": 'truefalse: Wahr/Falsch (question_data: {"correct_answer": true, "feedback_true": "...", "feedback_false": "..."})',
    "shortanswer": 'shortanswer: Kurzantwort (question_data: {"answers": [{"text": "Erwartete Antwort", "fraction": 100}]})',
    "numerical": 'numerical: Zahlenwert (question_data: {"answers": [{"value": 42, "tolerance": 0.1, "fraction": 100}]})',
    "matching": 'matching: Zuordnung (question_data: {"shuffle": true, "pairs": [{"question": "...", "answer": "..."}]})',
    "ordering": 'ordering: Reihenfolge (question_data: {"items": ["Erster", "Zweiter", "Dritter"]})',
    "cloze": 'cloze: Lückentext (question_data: {"gaps": [{"type": "shortanswer", "answers": [{"text": "...", "fraction": 100}]}]})',
    "essay": 'essay: Freitext (question_data: {"grader_info": "Erwartete Lösung und Kriterien"})',
    "webapp": 'webapp: Interaktive Web-App (question_data: {"app_description": "Beschreibung der App", "grader_info": "Bewertungskriterien"}) — nur wenn das Thema sich dafür eignet',
    "feynman": 'feynman: Erkläraufgabe — Schüler erklärt ein Konzept einem unwissenden KI-Kollegen im Chat (question_data: {"concept": "...", "context": "...", "max_turns": 10, "grader_info": "..."})',
    "scenario": 'scenario: Branching-Szenario — Schüler navigiert durch interaktive Entscheidungssituationen (question_data: {"scenario_description": "...", "context": "...", "max_decisions": 5, "grader_info": "..."})',
    "description": 'description: Nur Beschreibung/Info (question_data: {})',
    "photo": 'photo: Foto-Aufgabe — Schüler fotografiert reales Objekt/Ergebnis mit der Kamera (question_data: {"grader_info": "Was auf dem Foto zu sehen sein soll"}) — ideal für Werkstatt, Labor, Handwerk',
    "coding": 'coding: Programmieraufgabe (question_data: {"language": "javascript|python|sql|html|typescript", "starter_code": "Vorgabecode", "test_cases": [{"input": "funktionsaufruf(args)", "expected_output": "erwartetes Ergebnis", "description": "Testbeschreibung"}], "hidden_tests": false}) — IMMER mindestens 3 test_cases angeben! Bei SQL statt test_cases: {"sql_schema": "CREATE TABLE + INSERT ...", "sql_expected": [["spalte"], ["wert"]]}. Bei HTML: {"grader_info": "Bewertungskriterien"}',
}


def _build_task_prompt(allowed_types: list[str] | None = None, coding_language: str = "") -> str:
    """Build the task generation prompt, optionally filtered by allowed types."""
    if allowed_types:
        types = [t for t in allowed_types if t in _DOC_TYPE_DESCRIPTIONS]
    else:
        types = list(_DOC_TYPE_DESCRIPTIONS.keys())

    type_lines = "\n".join(f"- {_DOC_TYPE_DESCRIPTIONS[t]}" for t in types)

    coding_hint = ""
    if coding_language and "coding" in types:
        lang_labels = {"javascript": "JavaScript", "python": "Python", "sql": "SQL", "html": "HTML/CSS", "typescript": "TypeScript"}
        coding_hint = f"\n\nWICHTIG fuer Programmieraufgaben: Verwende die Sprache {lang_labels.get(coding_language, coding_language)} (language: \"{coding_language}\")."

    return f"""Analysiere dieses Dokument und erstelle daraus strukturierte Prüfungsaufgaben.

Verwende AUSSCHLIESSLICH diese Aufgabentypen:
{type_lines}

Wähle den Aufgabentyp basierend auf dem Inhalt:
- Wenn das Dokument bereits Aufgaben enthält, übernimm den passenden Typ
- Wenn es Lernstoff/Text ist, erstelle verschiedene Aufgabentypen dazu

WICHTIG für den Aufgabentext:
- Formuliere die Aufgaben in klarem, eigenständigem Deutsch
- Übernimm KEINE Formatierungsartefakte aus dem Quelldokument (z.B. Dateinamen wie "verzeichnisliste.txt", HTML/Markdown-Tags, Code-Blöcke die nur Formatierung sind)
- Der Aufgabentext muss für sich allein verständlich sein, ohne das Quelldokument

Antworte als JSON-Array:
[
  {{
    "title": "Aufgabentitel",
    "text": "Aufgabentext (Markdown)",
    "hint": "Optionaler Hinweis",
    "solution": "Ausführliche Musterlösung",
    "topic": "Themengebiet",
    "task_type": "...",
    "points": 1-5,
    "question_data": {{ ... }},
    "images": ["img_1"]
  }}
]

BILDER AUS DEM DOKUMENT: Jedes Bild im Dokument hat eine ID (img_1, img_2, ...).
Du kannst Bilder in Aufgaben einbinden, indem du die ID im "images"-Array der Aufgabe auflistest.
Die Bilder werden dann automatisch im Aufgabentext angezeigt.
- Verwende Originalbilder für: Fotos, Screenshots, handgezeichnete Skizzen, Abbildungen die man nicht als Text/Diagramm nachbauen kann
- Verwende STATTDESSEN Mermaid-Diagramme für: ER-Diagramme, Flowcharts, UML, Netzwerk-Topologien (sauberer, editierbar)
- Du kannst auch beides kombinieren: Ein Foto als Bild + ein Mermaid-Diagramm in derselben Aufgabe
- Wenn du ein Bild einbindest, schreibe im text-Feld an der gewünschten Stelle: {{{{img_1}}}} (wird durch das Bild ersetzt)

DIAGRAMME: Du kannst im Aufgabentext Mermaid-Diagramme verwenden! Sie werden automatisch gerendert.
Nutze ```mermaid Code-Blöcke für: ER-Diagramme, Flowcharts, Sequenzdiagramme, Klassendiagramme, Zustandsdiagramme.
Beispiel: "Beschreibe das folgende ER-Diagramm:\\n\\n```mermaid\\nerDiagram\\n    KUNDE ||--o{{ BESTELLUNG : bestellt\\n```"

WICHTIG: Jede Aufgabe MUSS eine ausführliche "solution" (Musterlösung) enthalten.
- Bei Multiple Choice: Erkläre warum die richtige Antwort korrekt ist
- Bei Kurzantwort/Numerisch: Gib die korrekte Antwort mit Erklärung
- Bei Essay/Freitext: Beschreibe was eine vollständige Antwort enthalten sollte
- Bei Zuordnung/Reihenfolge: Erkläre die korrekte Zuordnung/Reihenfolge
- Bei Programmierung: Erstelle IMMER mindestens 3 sinnvolle test_cases{coding_hint}"""

SYSTEM_PROMPT = """Du bist ein erfahrener Lehrer. Analysiere das hochgeladene Dokument und erstelle daraus strukturierte Prüfungsaufgaben in verschiedenen Formaten.
Antworte IMMER als valides JSON-Array. Keine zusätzliche Erklärung.
WICHTIG: Achte auf korrektes JSON-Escaping! Anführungszeichen innerhalb von String-Werten MÜSSEN escaped werden (\\"). Verwende keine unescapten " innerhalb von JSON-Strings."""


async def _post_process_tasks(tasks: list, image_map: dict[str, str] | None = None) -> None:
    """Validate tasks, embed referenced images, and generate app_html for webapp tasks."""
    from services.claude_service import generate_webapp

    for task in tasks:
        if not isinstance(task, dict):
            continue
        task.setdefault("title", "Unbenannte Aufgabe")
        task.setdefault("text", "")
        task.setdefault("hint", "")
        task.setdefault("solution", "")
        task.setdefault("topic", "")
        task.setdefault("task_type", "essay")
        task.setdefault("points", 1)
        task.setdefault("question_data", {})

        # Embed referenced images into task text as Markdown images
        if image_map:
            referenced_images = task.pop("images", []) or []
            for img_id in referenced_images:
                if img_id in image_map:
                    data_url = image_map[img_id]
                    placeholder = "{{" + img_id + "}}"
                    img_md = f"\n\n![{img_id}]({data_url})\n\n"
                    if placeholder in task["text"]:
                        task["text"] = task["text"].replace(placeholder, img_md)
                    else:
                        # No placeholder found — append image at end of text
                        task["text"] += img_md
                    # Also replace in solution if referenced
                    if placeholder in task.get("solution", ""):
                        task["solution"] = task["solution"].replace(placeholder, img_md)
            # Clean up any unreferenced placeholders
            task["text"] = re.sub(r'\{\{img_\d+\}\}', '', task["text"])

        # Generate app_html for webapp tasks
        if task["task_type"] == "webapp":
            qd = task["question_data"]
            desc = qd.get("app_description", "") or task.get("text", "")
            grader = qd.get("grader_info", "")
            try:
                app_html = await generate_webapp(desc, grader)
                qd["app_html"] = app_html
            except Exception:
                # Fallback to essay if webapp generation fails
                task["task_type"] = "essay"
                qd["grader_info"] = grader


def _extract_docx(file_path: str) -> tuple[list[dict], dict[str, str]]:
    """Extract text and images from a DOCX file.
    Returns (content_blocks_for_claude, image_map: {img_id: data_url})."""
    import docx

    doc = docx.Document(file_path)
    content = []
    image_map = {}

    # Extract full text
    paragraphs = []
    for para in doc.paragraphs:
        if para.text.strip():
            paragraphs.append(para.text)

    # Also extract text from tables
    for table in doc.tables:
        table_rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            table_rows.append(" | ".join(cells))
        if table_rows:
            paragraphs.append("\n".join(table_rows))

    if paragraphs:
        content.append({
            "type": "text",
            "text": "--- Dokumentinhalt ---\n\n" + "\n\n".join(paragraphs),
        })

    # Extract embedded images with IDs
    image_count = 0
    for rel in doc.part.rels.values():
        if "image" in rel.reltype:
            try:
                img_data = rel.target_part.blob
                # Detect image type
                if img_data[:8] == b'\x89PNG\r\n\x1a\n':
                    media_type = "image/png"
                elif img_data[:2] == b'\xff\xd8':
                    media_type = "image/jpeg"
                else:
                    continue  # Skip unsupported formats

                b64 = base64.b64encode(img_data).decode("utf-8")
                image_count += 1
                img_id = f"img_{image_count}"
                image_map[img_id] = f"data:{media_type};base64,{b64}"
                content.append({
                    "type": "text",
                    "text": f"--- Bild ID: {img_id} ---",
                })
                content.append({
                    "type": "image",
                    "source": {
                        "type": "base64",
                        "media_type": media_type,
                        "data": b64,
                    },
                })
            except Exception:
                continue

    return content, image_map


def _extract_pdf(file_path: str) -> tuple[list[dict], dict[str, str]]:
    """Render PDF pages as images + extract embedded images.
    Returns (content_blocks_for_claude, image_map: {img_id: data_url})."""
    import fitz

    doc = fitz.open(file_path)
    content = []
    image_map = {}

    # Limit to 20 pages
    page_count = min(len(doc), 20)
    image_count = 0

    for i in range(page_count):
        page = doc[i]
        pix = page.get_pixmap(dpi=200)
        img_bytes = pix.tobytes("png")
        b64 = base64.b64encode(img_bytes).decode("utf-8")

        content.append({
            "type": "text",
            "text": f"--- Seite {i + 1} ---",
        })
        content.append({
            "type": "image",
            "source": {
                "type": "base64",
                "media_type": "image/png",
                "data": b64,
            },
        })

        # Extract individual images from this page
        for img_info in page.get_images(full=True):
            try:
                xref = img_info[0]
                base_image = doc.extract_image(xref)
                if not base_image or base_image["width"] < 50 or base_image["height"] < 50:
                    continue  # Skip tiny images (icons, bullets)
                img_ext = base_image["ext"]
                if img_ext not in ("png", "jpeg", "jpg"):
                    continue
                media_type = f"image/{'jpeg' if img_ext == 'jpg' else img_ext}"
                img_b64 = base64.b64encode(base_image["image"]).decode("utf-8")
                image_count += 1
                img_id = f"img_{image_count}"
                image_map[img_id] = f"data:{media_type};base64,{img_b64}"
                content.append({
                    "type": "text",
                    "text": f"--- Einzelbild auf Seite {i+1}, ID: {img_id} ---",
                })
                content.append({
                    "type": "image",
                    "source": {
                        "type": "base64",
                        "media_type": media_type,
                        "data": img_b64,
                    },
                })
            except Exception:
                continue

    doc.close()
    return content, image_map


async def import_document(file_path: str, original_filename: str, allowed_types: list[str] | None = None, coding_language: str = "") -> list[dict]:
    """
    Import tasks from a document (PDF or DOCX).
    - DOCX: Extracts text + images directly via python-docx
    - PDF: Renders pages as images via PyMuPDF
    Sends content to Claude for task extraction.
    Returns list of task dicts with question_data.
    """
    ext = original_filename.rsplit(".", 1)[-1].lower() if "." in original_filename else ""

    if ext == "docx":
        content, image_map = _extract_docx(file_path)
    else:
        content, image_map = _extract_pdf(file_path)

    if not content:
        raise ValueError("Keine Inhalte im Dokument gefunden.")

    # Append task generation prompt
    content.append({"type": "text", "text": _build_task_prompt(allowed_types, coding_language)})

    async with _semaphore:
        response = await asyncio.to_thread(
            _get_client().messages.create,
            model=CLAUDE_MODEL,
            max_tokens=64000,
            system=SYSTEM_PROMPT,
            messages=[{"role": "user", "content": content}],
        )

    tasks = _parse_json_response(response.content[0].text)

    # Validate, embed images, and post-process
    await _post_process_tasks(tasks, image_map)

    return tasks


async def import_document_with_instructions(
    file_path: str, original_filename: str, instructions: str, allowed_types: list[str] | None = None
) -> list[dict]:
    """
    Like import_document, but with custom teacher instructions for task generation.
    The instructions guide what kind of tasks to create.
    """
    ext = original_filename.rsplit(".", 1)[-1].lower() if "." in original_filename else ""

    if ext == "docx":
        content, image_map = _extract_docx(file_path)
    else:
        content, image_map = _extract_pdf(file_path)

    if not content:
        raise ValueError("Keine Inhalte im Dokument gefunden.")

    # Extract coding language from instructions if present
    import re as _re
    lang_match = _re.search(r'language:\s*"(\w+)"', instructions)
    _coding_lang = lang_match.group(1) if lang_match else ""

    # Build a customized prompt with teacher instructions
    custom_prompt = _build_task_prompt(allowed_types, _coding_lang)
    if instructions.strip():
        custom_prompt += f"\n\nWICHTIG — Anweisungen des Lehrers:\n{instructions.strip()}\n\nHalte dich unbedingt an diese Anweisungen bei der Erstellung der Aufgaben."

    content.append({"type": "text", "text": custom_prompt})

    async with _semaphore:
        response = await asyncio.to_thread(
            _get_client().messages.create,
            model=CLAUDE_MODEL,
            max_tokens=64000,
            system=SYSTEM_PROMPT,
            messages=[{"role": "user", "content": content}],
        )

    tasks = _parse_json_response(response.content[0].text)

    # Validate, embed images, and post-process
    await _post_process_tasks(tasks, image_map)

    return tasks
